# ===============================================================
# GENERADOR DE INFORMES DE AUDITORÍA EN FORMATO DOCX - V2
# Versión mejorada con mejor manejo de errores
# Conforme a RT 7, RT 37 y NIAs
# ===============================================================

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import pandas as pd
import traceback
import sys

class GeneradorInformeAuditoria:
    """Genera informes profesionales de auditoría en formato DOCX"""
    
    def __init__(self, empresa_nombre, empresa_cuit, fecha_auditoria):
        try:
            self.empresa_nombre = empresa_nombre
            self.empresa_cuit = empresa_cuit
            self.fecha_auditoria = fecha_auditoria
            self.doc = Document()
            self._configurar_estilos()
            print(f"✓ Generador inicializado correctamente")
        except Exception as e:
            print(f"✗ Error al inicializar generador: {e}")
            raise
    
    def _configurar_estilos(self):
        """Configura los estilos del documento"""
        try:
            styles = self.doc.styles
            
            # Título principal (solo si no existe)
            if 'Titulo Principal' not in [s.name for s in styles]:
                titulo_style = styles.add_style('Titulo Principal', WD_STYLE_TYPE.PARAGRAPH)
                titulo_format = titulo_style.font
                titulo_format.name = 'Arial'
                titulo_format.size = Pt(16)
                titulo_format.bold = True
                titulo_format.color.rgb = RGBColor(0, 51, 102)
            print(f"✓ Estilos configurados")
        except Exception as e:
            print(f"✗ Error al configurar estilos: {e}")
            # No es crítico, continuamos
    
    def agregar_portada(self):
        """Agrega la portada del informe"""
        try:
            # Título principal
            titulo = self.doc.add_paragraph()
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = titulo.add_run("INFORME DE AUDITORÍA\nSOBRE ACTIVOS CORRIENTES")
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            
            self.doc.add_paragraph()  # Espacio
            
            # Datos de la empresa
            datos = self.doc.add_paragraph()
            datos.alignment = WD_ALIGN_PARAGRAPH.CENTER
            datos_text = f"""
Empresa Auditada: {self.empresa_nombre}
CUIT: {self.empresa_cuit}
Período Auditado: {self.fecha_auditoria.strftime('%m/%Y')}
Fecha del Informe: {datetime.now().strftime('%d de %B de %Y')}

Normas Aplicadas: RT 7, RT 37, NIAs
            """
            datos.add_run(datos_text)
            
            self.doc.add_page_break()
            print(f"✓ Portada agregada")
        except Exception as e:
            print(f"✗ Error al agregar portada: {e}")
            raise
    
    def agregar_seccion_identificacion(self):
        """Agrega la sección de identificación del ente"""
        try:
            heading = self.doc.add_heading('I. IDENTIFICACIÓN DEL ENTE Y PERÍODO AUDITADO', level=1)
            
            texto = f"""Hemos auditado los activos corrientes de {self.empresa_nombre}, CUIT {self.empresa_cuit}, correspondientes al período finalizado el {self.fecha_auditoria.strftime('%d de %B de %Y')}. La auditoría se realizó conforme a las Normas Internacionales de Auditoría (NIAs) y las Resoluciones Técnicas 7 y 37 de FACPCE."""
            
            p = self.doc.add_paragraph(texto)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            print(f"✓ Sección identificación agregada")
        except Exception as e:
            print(f"✗ Error en identificación: {e}")
            raise
    
    def agregar_seccion_alcance(self):
        """Agrega la sección de alcance del trabajo"""
        try:
            self.doc.add_heading('II. ALCANCE DEL TRABAJO', level=1)
            self.doc.add_heading('2.1. Procedimientos Aplicados (Según RT 7)', level=2)
            
            procedimientos = [
                "Confirmaciones externas de saldos",
                "Inspección de documentación respaldatoria",
                "Pruebas de corte de operaciones",
                "Análisis de antigüedad de saldos",
                "Revisión de conciliaciones bancarias",
                "Verificación de cálculos de intereses",
                "Evaluación de controles internos",
                "Pruebas sustantivas de transacciones",
                "Análisis de eventos posteriores",
                "Aplicación de técnicas analíticas con Machine Learning",
                "Evaluación de recuperabilidad de activos",
                "Verificación de valuación conforme RT 17 y RT 31"
            ]
            
            for proc in procedimientos:
                self.doc.add_paragraph(proc, style='List Bullet')
            
            print(f"✓ Sección alcance agregada")
        except Exception as e:
            print(f"✗ Error en alcance: {e}")
            raise
    
    def agregar_resumen_hallazgos(self, resumen_df):
        """Agrega el resumen de hallazgos con tabla"""
        try:
            self.doc.add_heading('III. RESUMEN DE HALLAZGOS', level=1)
            
            # Crear tabla
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Encabezados
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'RUBRO'
            hdr_cells[1].text = 'CANTIDAD'
            hdr_cells[2].text = 'SALDO ($)'
            hdr_cells[3].text = '% TOTAL'
            
            # Hacer encabezados en negrita
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Calcular total para porcentajes
            total_general = resumen_df[resumen_df['Rubro'] == 'TOTAL ACTIVOS CORRIENTES']['Saldo ($)'].iloc[0]
            
            # Agregar datos
            for _, row in resumen_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['Rubro'])
                row_cells[1].text = f"{int(row['Cantidad'])}"
                row_cells[2].text = f"{row['Saldo ($)']:,.2f}"
                
                if row['Rubro'] != 'TOTAL ACTIVOS CORRIENTES':
                    porcentaje = (row['Saldo ($)'] / total_general * 100) if total_general > 0 else 0
                    row_cells[3].text = f"{porcentaje:.1f}%"
                else:
                    row_cells[3].text = "100.0%"
                    # Poner última fila en negrita
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            
            print(f"✓ Resumen de hallazgos agregado")
        except Exception as e:
            print(f"✗ Error en resumen: {e}")
            raise
    
    def agregar_hallazgos_especificos(self, data_dict):
        """Agrega hallazgos específicos por rubro"""
        try:
            self.doc.add_heading('IV. HALLAZGOS ESPECÍFICOS POR RUBRO', level=1)
            
            for rubro, df in data_dict.items():
                self.doc.add_heading(f'4.{list(data_dict.keys()).index(rubro)+1}. {rubro}', level=2)
                
                # Estadísticas
                total_items = len(df)
                anomalias = len(df[df.get('resultado_if', pd.Series()) == 'Anómalo'])
                alertas = len(df[df.get('alerta', pd.Series()).notna()])
                
                self.doc.add_paragraph(f"Total de items auditados: {total_items}")
                self.doc.add_paragraph(f"Anomalías detectadas (Machine Learning): {anomalias}")
                self.doc.add_paragraph(f"Alertas por reglas de negocio: {alertas}")
                
                # Observaciones
                if anomalias > 0 or alertas > 0:
                    self.doc.add_paragraph(
                        f"⚠️ Observación: Se detectaron {anomalias + alertas} items con características atípicas que requieren análisis adicional.",
                        style='Intense Quote'
                    )
                    
                    # Recomendación según RT 7
                    self.doc.add_paragraph("Recomendación (RT 7):", style='Heading 3')
                    
                    if rubro == 'Caja y Bancos':
                        rec = "Implementar sistema de alertas tempranas y realizar arqueos sorpresivos."
                    elif rubro == 'Cuentas a Cobrar':
                        rec = "Circularizar clientes y evaluar la recuperabilidad de saldos vencidos."
                    elif rubro == 'Inventarios':
                        rec = "Realizar recuentos físicos y evaluar obsolescencia según RT 31."
                    elif 'Inversiones' in rubro:
                        rec = "Obtener confirmaciones de custodios y verificar valuaciones de mercado."
                    else:
                        rec = "Revisar documentación respaldatoria y verificar devengamiento correcto."
                    
                    self.doc.add_paragraph(rec)
                else:
                    self.doc.add_paragraph(
                        "✓ No se detectaron observaciones significativas en este rubro.",
                        style='Quote'
                    )
                
                self.doc.add_paragraph()  # Espacio
            
            print(f"✓ Hallazgos específicos agregados")
        except Exception as e:
            print(f"✗ Error en hallazgos específicos: {e}")
            raise
    
    def agregar_opinion_profesional(self, total_activos, total_anomalias, total_items):
        """Agrega la opinión profesional del auditor"""
        try:
            self.doc.add_heading('V. OPINION PROFESIONAL', level=1)
            
            # Determinar tipo de opinión
            tasa_anomalias = (total_anomalias / total_items * 100) if total_items > 0 else 0
            
            if tasa_anomalias < 5:
                opinion_tipo = "sin salvedades"
                texto_opinion = f"""En nuestra opinión, basada en la auditoría realizada conforme a las NIAs y RT 7 y 37, los activos corrientes de {self.empresa_nombre} al {self.fecha_auditoria.strftime('%d de %B de %Y')}, por un total de ${total_activos:,.2f}, se presentan razonablemente en todos sus aspectos significativos, de conformidad con las normas contables profesionales vigentes en Argentina."""
            elif tasa_anomalias < 10:
                opinion_tipo = "con salvedades"
                texto_opinion = f"""En nuestra opinión, excepto por los efectos de los asuntos mencionados en la sección de Hallazgos Específicos, los activos corrientes de {self.empresa_nombre} al {self.fecha_auditoria.strftime('%d de %B de %Y')}, por un total de ${total_activos:,.2f}, se presentan razonablemente en todos sus aspectos significativos."""
            else:
                opinion_tipo = "adversa"
                texto_opinion = f"""En nuestra opinión, debido a la importancia de los asuntos mencionados en la sección de Hallazgos Específicos, los activos corrientes de {self.empresa_nombre} al {self.fecha_auditoria.strftime('%d de %B de %Y')} no se presentan razonablemente de conformidad con las normas contables profesionales."""
            
            p = self.doc.add_paragraph(texto_opinion)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            self.doc.add_paragraph()
            
            # Párrafo de énfasis sobre anomalías si aplica
            if total_anomalias > 0:
                enfasis = self.doc.add_paragraph()
                enfasis.add_run("Párrafo de Énfasis: ").bold = True
                enfasis.add_run(
                    f"Llamamos la atención sobre la existencia de {total_anomalias} ítems detectados con características atípicas mediante técnicas de análisis de datos y Machine Learning, los cuales requieren investigación adicional por parte de la administración."
                )
                enfasis.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            print(f"✓ Opinión profesional agregada ({opinion_tipo})")
        except Exception as e:
            print(f"✗ Error en opinión: {e}")
            raise
    
    def agregar_firmas(self):
        """Agrega sección de firmas"""
        try:
            self.doc.add_paragraph()
            self.doc.add_paragraph()
            
            # Tabla para firmas
            table = self.doc.add_table(rows=3, cols=2)
            
            # Líneas de firma
            table.cell(0, 0).text = "_" * 35
            table.cell(0, 1).text = "_" * 35
            
            # Nombres
            table.cell(1, 0).text = "Contador Público"
            table.cell(1, 1).text = "Socio Director"
            
            # Matrícula
            table.cell(2, 0).text = "CPCECABA T° XXX F° XXX"
            table.cell(2, 1).text = "Estudio Contable"
            
            # Centrar texto
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self.doc.add_paragraph()
            lugar_fecha = self.doc.add_paragraph(
                f"Buenos Aires, {datetime.now().strftime('%d de %B de %Y')}"
            )
            lugar_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"✓ Firmas agregadas")
        except Exception as e:
            print(f"✗ Error en firmas: {e}")
            raise
    
    def agregar_anexos(self, data_dict):
        """Agrega anexos con detalles de items con observaciones"""
        try:
            self.doc.add_page_break()
            self.doc.add_heading('ANEXO I - DETALLE DE ITEMS CON OBSERVACIONES', level=1)
            
            items_anexo = 0
            for rubro, df in data_dict.items():
                df_problemas = df[(df.get('resultado_if', pd.Series()) == 'Anómalo') | 
                                 (df.get('alerta', pd.Series()).notna())]
                
                if not df_problemas.empty:
                    items_anexo += 1
                    self.doc.add_heading(rubro, level=2)
                    self.doc.add_paragraph(f"Total de items con observaciones: {len(df_problemas)}")
                    
                    # Crear tabla con primeras 10 observaciones
                    if len(df_problemas) > 0:
                        # Seleccionar columnas relevantes
                        cols_mostrar = [col for col in df_problemas.columns 
                                      if col not in ['anomaly_if', 'resultado_if', 'alerta']][:5]
                        
                        df_muestra = df_problemas[cols_mostrar].head(10)
                        
                        # Crear tabla
                        table = self.doc.add_table(rows=1, cols=len(cols_mostrar))
                        table.style = 'Light List Accent 1'
                        
                        # Encabezados
                        hdr_cells = table.rows[0].cells
                        for i, col in enumerate(cols_mostrar):
                            hdr_cells[i].text = str(col)
                            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                        
                        # Datos
                        for _, row in df_muestra.iterrows():
                            row_cells = table.add_row().cells
                            for i, col in enumerate(cols_mostrar):
                                row_cells[i].text = str(row[col])[:50]  # Limitar longitud
                    
                    self.doc.add_paragraph()
            
            print(f"✓ Anexos agregados ({items_anexo} rubros con observaciones)")
        except Exception as e:
            print(f"✗ Error en anexos: {e}")
            # No es crítico
            print(f"  Continuando sin anexos...")
    
    def generar_informe(self, resumen_df, data_dict, ruta_salida):
        """Genera el informe completo"""
        try:
            print(f"\n{'='*60}")
            print(f"GENERANDO INFORME DE AUDITORÍA")
            print(f"{'='*60}")
            print(f"Empresa: {self.empresa_nombre}")
            print(f"Ruta: {ruta_salida}")
            print(f"{'='*60}\n")
            
            # Agregar todas las secciones
            self.agregar_portada()
            self.agregar_seccion_identificacion()
            self.agregar_seccion_alcance()
            self.agregar_resumen_hallazgos(resumen_df)
            self.agregar_hallazgos_especificos(data_dict)
            
            # Calcular métricas para opinión
            total_activos = resumen_df[resumen_df['Rubro'] == 'TOTAL ACTIVOS CORRIENTES']['Saldo ($)'].iloc[0]
            total_anomalias = resumen_df[resumen_df['Rubro'] == 'TOTAL ACTIVOS CORRIENTES']['Anomalías'].iloc[0]
            total_items = resumen_df[resumen_df['Rubro'] == 'TOTAL ACTIVOS CORRIENTES']['Cantidad'].iloc[0]
            
            self.agregar_opinion_profesional(total_activos, total_anomalias, total_items)
            self.agregar_firmas()
            self.agregar_anexos(data_dict)
            
            # Guardar documento
            self.doc.save(ruta_salida)
            
            print(f"\n{'='*60}")
            print(f"✅ INFORME GENERADO EXITOSAMENTE")
            print(f"{'='*60}\n")
            
            return ruta_salida
            
        except Exception as e:
            print(f"\n{'='*60}")
            print(f"❌ ERROR AL GENERAR INFORME")
            print(f"{'='*60}")
            print(f"Error: {str(e)}")
            print(f"Tipo: {type(e).__name__}")
            print(f"\nTraceback completo:")
            traceback.print_exc()
            print(f"{'='*60}\n")
            raise
